import docx
from copy import deepcopy
from datetime import date

from test_xlsx import find_items


CUR_DATE = date.today().strftime('%d.%m.%Y')
DOCUMENT = docx.Document('source\\blank.docx')
TABLES = DOCUMENT.tables
PARAGRAPHS_POSITION = {
    'id': [DOCUMENT, 3, 2,],
    'date': [DOCUMENT, 3, 5,],
    'room': [DOCUMENT, 5, 2,]
}


def generate_table(len_of_data_xl):
    tables_position = {
        'department': [TABLES, 0, 0, 1],
        'name_pc': [TABLES, 0, 1, 1],
        'model_pc': [TABLES, 1, 1, 1],
        'serialnum_pc': [TABLES, 1, 1, 2],
        'bu_pc': [TABLES, 1, 1, 3],
        'model_monitor_one': [TABLES, 1, 2, 1],
        'serialnum_monitor_one': [TABLES, 1, 2, 2],
        'bu_monitor_one': [TABLES, 1, 2, 3],
    }
    tables_len_three = {
        'model_monitor_two': [TABLES, 1, 3, 1],
        'serialnum_monitor_two': [TABLES, 1, 3, 2],
        'bu_monitor_two': [TABLES, 1, 3, 3],
    }
    tables_len_four = {
        'model_monitor_three': [TABLES, 1, 4, 1],
        'serialnum_monitor_three': [TABLES, 1, 4, 2],
        'bu_monitor_three': [TABLES, 1, 4, 3],
    }
    if len_of_data_xl == 2:
        for item in tables_len_three.values():
            change_table_text(*item, text='')
        for item in tables_len_four.values():
            change_table_text(*item, text='')
        change_table_text(TABLES, 1, 3, 0, '')
        change_table_text(TABLES, 1, 4, 0, '')
    elif len_of_data_xl == 3:
        tables_position.update(tables_len_three)
        for item in tables_len_four.values():
            change_table_text(*item, text='')
        change_table_text(TABLES, 1, 4, 0, '')
    elif len_of_data_xl == 4:
        tables_position.update(tables_len_three)
        tables_position.update(tables_len_four)
    return tables_position


def change_paragraph_text(doc_obj, num_par, num_run, text):
    p = doc_obj.paragraphs[num_par]
    run = p.runs[num_run]
    run.text = text


def change_table_text(table_obj, num_table, num_rows, num_cells, text):
    obj = table_obj[num_table].rows[num_rows].cells[num_cells]
    change_paragraph_text(obj, 0, 0, text)


def general_bu(data, pc):
    pc_bu = pc[5]
    for line in data:
        if line[5] == pc_bu:
            change_table_text(TABLES, 0, 2, 1, pc_bu)
            return True
    change_table_text(TABLES, 0, 2, 1, '')
    return False


def create_doc(data, tables_position):
    data_from_xl = deepcopy(data)
    len_data = len(data_from_xl)
    for i in range(len_data):
        if data_from_xl[i][2] == 'системный блок':
            pc = data_from_xl.pop(i)
            break
    general_bu(data_from_xl, pc)
    for i in range(len(data_from_xl)):
        if data_from_xl[i][5] == pc[5]:
            data_from_xl[i], data_from_xl[0] = data_from_xl[0], data_from_xl[i]
            break
    tables_text = {
        'department': pc[0],
        'name_pc': pc[4],
        'model_pc': pc[3],
        'serialnum_pc': pc[6],
        'bu_pc': pc[5],
        'model_monitor_one': data_from_xl[0][3],
        'serialnum_monitor_one': data_from_xl[0][6],
        'bu_monitor_one': data_from_xl[0][5],
    }
    if len(data_from_xl) > 1:
        monitor_two = {
            'model_monitor_two': data_from_xl[1][3],
            'serialnum_monitor_two': data_from_xl[1][6],
            'bu_monitor_two': data_from_xl[1][5],
        }
        tables_text.update(monitor_two)
    if len(data_from_xl) == 3:
        monitor_three = {
            'model_monitor_three': data_from_xl[2][3],
            'serialnum_monitor_three': data_from_xl[2][6],
            'bu_monitor_three': data_from_xl[2][5],
        }
        tables_text.update(monitor_three)
    for key, value in tables_text.items():
        change_table_text(*tables_position[key], value)
    change_paragraph_text(*PARAGRAPHS_POSITION['id'], text=pc[7])
    change_paragraph_text(*PARAGRAPHS_POSITION['date'], text=CUR_DATE)
    change_paragraph_text(*PARAGRAPHS_POSITION['room'], text=pc[1])
    print('Текст успешно перенесен')


def main():
    while True:
        print('Введите имя компьютера: ')
        find = input()
        data_from_xl = find_items(find)
        tables_position = generate_table(len(data_from_xl))
        create_doc(data_from_xl, tables_position)
        print(f'Паспорт с именем {find}.docx в папке \\source\\passports сохранен\n\n')
        DOCUMENT.save(f'source\\passports\\{find}.docx')


if __name__ == '__main__':
    main()
