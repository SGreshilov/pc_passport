import docx


doc = docx.Document('source\\blank.docx')
paragraphs = doc.paragraphs
for i in range(len(paragraphs)):
    print(f'строка {i}: {paragraphs[i].text}')
tables = doc.tables
print(tables[1].rows[4].cells[2].paragraphs[0].runs)